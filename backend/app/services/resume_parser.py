"""
Module B — file validation + text extraction for PDF/DOCX resumes.
"""
import io
import logging
import re

import pdfplumber
from docx import Document

from app.config import get_settings
from app.core.exceptions import ExtractionError, FileValidationError

logger = logging.getLogger("skillforge.resume_parser")

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    # Some browsers/clients send this generic type for .docx
    "application/octet-stream",
}


def validate_upload(filename: str | None, content_type: str | None, content: bytes) -> None:
    if not filename:
        raise FileValidationError("No filename was provided with the upload.")

    ext = _extension(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise FileValidationError(
            f"Unsupported file type '{ext or 'unknown'}'. Please upload a PDF or DOCX resume.",
            detail={"allowed": sorted(ALLOWED_EXTENSIONS)},
        )

    if not content:
        raise FileValidationError("The uploaded file is empty.")

    settings = get_settings()
    if len(content) > settings.max_resume_file_size_bytes:
        raise FileValidationError(
            f"File is too large ({len(content) / (1024 * 1024):.1f} MB). "
            f"Max size is {settings.MAX_RESUME_FILE_SIZE_MB} MB.",
        )

    if content_type and content_type not in ALLOWED_CONTENT_TYPES and ext == ".pdf" and not content.startswith(b"%PDF"):
        raise FileValidationError("This file doesn't look like a valid PDF.")


def _extension(filename: str) -> str:
    idx = filename.rfind(".")
    return filename[idx:].lower() if idx != -1 else ""


def extract_text(filename: str, content: bytes) -> str:
    ext = _extension(filename)
    if ext == ".pdf":
        text = _extract_pdf_text(content)
    elif ext == ".docx":
        text = _extract_docx_text(content)
    else:
        raise FileValidationError(f"Unsupported file type '{ext}'.")

    normalized = normalize_text(text)
    if not normalized:
        raise ExtractionError(
            "No readable text could be found in this resume. It may be a scanned image, "
            "password-protected, or empty — try uploading a text-based PDF or DOCX instead."
        )
    return normalized


def _extract_pdf_text(content: bytes) -> str:
    try:
        chunks: list[str] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                chunks.append(page_text)
        return "\n".join(chunks)
    except Exception as exc:  # noqa: BLE001
        logger.warning("PDF extraction failed: %s", exc)
        raise ExtractionError(
            "This PDF couldn't be read. It may be corrupted, encrypted, or a scanned image "
            "without a text layer."
        ) from exc


def _extract_docx_text(content: bytes) -> str:
    try:
        document = Document(io.BytesIO(content))
        parts: list[str] = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        logger.warning("DOCX extraction failed: %s", exc)
        raise ExtractionError(
            "This DOCX file couldn't be read. It may be corrupted or not a valid Word document."
        ) from exc


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
